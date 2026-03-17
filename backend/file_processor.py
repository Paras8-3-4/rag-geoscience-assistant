import os
import sys
sys.path.insert(0, os.path.dirname(__file__))

import tempfile
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from config import CHUNK_SIZE, CHUNK_OVERLAP

MAX_FILES = 10
MAX_TOTAL_SIZE_MB = 150
ALLOWED_EXTENSIONS = [".pdf", ".txt", ".docx", ".xlsx", ".xls"]


def validate_uploads(uploaded_files: list) -> tuple[bool, str]:
    if len(uploaded_files) > MAX_FILES:
        return False, f"Too many files. Maximum allowed is {MAX_FILES} files."

    total_size = sum(f.size for f in uploaded_files)
    total_size_mb = total_size / (1024 * 1024)
    if total_size_mb > MAX_TOTAL_SIZE_MB:
        return False, f"Total size {total_size_mb:.1f}MB exceeds {MAX_TOTAL_SIZE_MB}MB limit."

    for f in uploaded_files:
        ext = os.path.splitext(f.filename)[1].lower()
        if ext not in ALLOWED_EXTENSIONS:
            return False, f"File '{f.filename}' has unsupported format. Allowed: PDF, TXT, DOCX, XLSX."

    return True, "OK"


def extract_text_from_file(file_bytes: bytes, filename: str) -> list[Document]:
    ext = os.path.splitext(filename)[1].lower()
    documents = []

    try:
        if ext == ".pdf":
            documents = _process_pdf(file_bytes, filename)

        elif ext == ".txt":
            documents = _process_txt(file_bytes, filename)

        elif ext == ".docx":
            documents = _process_docx(file_bytes, filename)

        elif ext in [".xlsx", ".xls"]:
            documents = _process_excel(file_bytes, filename)

    except Exception as e:
        print(f"Error processing {filename}: {e}")
        return []

    return documents


def _process_pdf(file_bytes: bytes, filename: str) -> list[Document]:
    from pypdf import PdfReader
    import io
    reader = PdfReader(io.BytesIO(file_bytes))
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": filename, "page": i + 1}
            ))
    return docs


def _process_txt(file_bytes: bytes, filename: str) -> list[Document]:
    text = file_bytes.decode("utf-8", errors="ignore")
    if text.strip():
        return [Document(
            page_content=text,
            metadata={"source": filename, "page": 1}
        )]
    return []


def _process_docx(file_bytes: bytes, filename: str) -> list[Document]:
    from docx import Document as DocxDocument
    import io
    doc = DocxDocument(io.BytesIO(file_bytes))
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    if full_text.strip():
        return [Document(
            page_content=full_text,
            metadata={"source": filename, "page": 1}
        )]
    return []


def _process_excel(file_bytes: bytes, filename: str) -> list[Document]:
    import openpyxl
    import io
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    docs = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_text = []
        for row in ws.iter_rows(values_only=True):
            row_str = " | ".join([str(cell) for cell in row if cell is not None])
            if row_str.strip():
                rows_text.append(row_str)
        if rows_text:
            docs.append(Document(
                page_content="\n".join(rows_text),
                metadata={"source": filename, "page": sheet_name}
            ))
    return docs


def chunk_documents(documents: list[Document]) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP
    )
    return splitter.split_documents(documents)